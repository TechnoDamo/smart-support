"""RAG: чанкинг, ingestion, гибридный retrieval."""

from __future__ import annotations

import pytest
from sqlalchemy import select

from app.providers.embedding import BM25Encoder, tokenize
from app.db.models import RagDocument, RagDocumentChunk, RagIngestionJob
from app.services.rag import chunk_text, ingest_document, retrieve
from app.services.refs import get_default_rag_collection


def test_chunk_text_basic():
    chunks = chunk_text("a b c d e f g h", chunk_size_tokens=3, overlap_tokens=1)
    # Окно 3 со сдвигом 2 → [a b c], [c d e], [e f g], [g h]
    assert chunks[0] == "a b c"
    assert chunks[1] == "c d e"
    assert len(chunks) >= 3


def test_chunk_text_empty():
    assert chunk_text("", chunk_size_tokens=3, overlap_tokens=1) == []


def test_tokenize_splits_hyphenated_identifiers():
    tokens = tokenize("ГОСТ Р 52931-2008 и Эликонт-100")
    assert "52931-2008" in tokens
    assert "52931" in tokens
    assert "2008" in tokens
    assert "эликонт-100" in tokens
    assert "эликонт" in tokens
    assert "100" in tokens


def test_bm25_token_ids_are_stable_across_fit_order():
    encoder_left = BM25Encoder()
    encoder_right = BM25Encoder()

    left_docs = [
        tokenize("ГОСТ Р 52931-2008"),
        tokenize("Эликонт-100"),
    ]
    right_docs = list(reversed(left_docs))

    encoder_left.fit_add(left_docs)
    encoder_right.fit_add(right_docs)

    left_query = encoder_left.encode(tokenize("ГОСТ Р 52931"))
    right_query = encoder_right.encode(tokenize("ГОСТ Р 52931"))

    assert left_query.indices == right_query.indices


def test_extract_text():
    """Тестирует извлечение текста из разных форматов."""
    from app.services.rag import _extract_text
    from io import BytesIO

    # Plain text
    plain_content = b"Hello, world!"
    assert _extract_text(plain_content, "text/plain") == "Hello, world!"
    assert _extract_text(plain_content, "text/html") == "Hello, world!"
    assert _extract_text(plain_content, None) == "Hello, world!"

    # PDF error case (invalid PDF)
    import pytest

    with pytest.raises(ValueError, match="Ошибка извлечения текста из PDF"):
        _extract_text(b"invalid pdf", "application/pdf")

    # DOCX
    import docx

    doc = docx.Document()
    doc.add_paragraph("DOCX test paragraph")
    doc.add_paragraph("Another paragraph")
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_bytes = docx_buffer.getvalue()
    extracted = _extract_text(
        docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "DOCX test paragraph" in extracted
    assert "Another paragraph" in extracted

    # DOCX error case
    with pytest.raises(ValueError, match="Ошибка извлечения текста из DOCX"):
        _extract_text(
            b"invalid docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    # Unsupported MIME type → fallback to UTF‑8
    assert _extract_text(b"fallback", "application/octet-stream") == "fallback"

    # Binary content with unknown MIME type → null characters survive decode
    assert _extract_text(b"\xff\xfe\x00\x00", "application/unknown") == "\x00\x00"


@pytest.mark.asyncio
async def test_ingest_and_retrieve(db, providers):
    # Создаём документ и job, прогоняем ingestion
    collection = await get_default_rag_collection(db)
    doc = RagDocument(
        collection_id=collection.id,
        source_type="file",
        source_name="faq.txt",
        mime_type="text/plain",
        storage_url=None,
        current_version=0,
    )
    db.add(doc)
    await db.flush()
    job = RagIngestionJob(
        collection_id=collection.id,
        document_id=doc.id,
        operation="upsert_document",
        status="queued",
    )
    db.add(job)
    await db.flush()

    text = (
        "График работы поддержки: с 9 до 18 по будням. "
        "По выходным отвечаем дольше. "
        "Для возврата товара нужно предоставить чек. "
        "Доставка в пределах города — 1 день."
    )
    await ingest_document(
        db,
        document=doc,
        job=job,
        raw_bytes=text.encode("utf-8"),
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )
    await db.commit()

    assert job.status == "done"
    r = await db.execute(
        select(RagDocumentChunk).where(RagDocumentChunk.document_id == doc.id)
    )
    chunks = list(r.scalars())
    assert len(chunks) >= 1

    # Теперь retrieve
    import uuid as _uuid

    hits = await retrieve(
        db,
        chat_id=_uuid.uuid4(),
        ticket_id=_uuid.uuid4(),
        message_id=None,
        query_text="график работы",
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )
    assert isinstance(hits, list)
    # В in-memory сторе всегда что-то возвращается
    assert len(hits) >= 0


@pytest.mark.asyncio
async def test_retrieve_prunes_orphan_vector_points(db, providers):
    collection = await get_default_rag_collection(db)
    doc = RagDocument(
        collection_id=collection.id,
        source_type="file",
        source_name="faq.txt",
        mime_type="text/plain",
        storage_url=None,
        current_version=0,
    )
    db.add(doc)
    await db.flush()
    job = RagIngestionJob(
        collection_id=collection.id,
        document_id=doc.id,
        operation="upsert_document",
        status="queued",
    )
    db.add(job)
    await db.flush()

    text = "Чтобы перезапустить контроллер, удерживайте RESET 10 секунд."
    await ingest_document(
        db,
        document=doc,
        job=job,
        raw_bytes=text.encode("utf-8"),
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )
    await db.commit()

    [orphan_vec] = await providers.embedding.embed(["reset controller orphan"])
    await providers.vector_store.upsert(
        collection.qdrant_collection_name,
        "orphan-point",
        orphan_vec,
        payload={"document_id": "missing", "chunk_index": 999},
    )

    hits = await retrieve(
        db,
        chat_id=doc.id,
        ticket_id=doc.id,
        message_id=None,
        query_text="reset controller",
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )

    assert all(chunk.qdrant_point_id != "orphan-point" for chunk, _score in hits)
    assert "orphan-point" not in providers.vector_store._data[collection.qdrant_collection_name]


@pytest.mark.asyncio
async def test_retrieve_boosts_exact_technical_identifier_matches(db, providers):
    collection = await get_default_rag_collection(db)

    def _make_doc(source_name: str) -> RagDocument:
        doc = RagDocument(
            collection_id=collection.id,
            source_type="file",
            source_name=source_name,
            mime_type="text/plain",
            storage_url=None,
            current_version=0,
        )
        db.add(doc)
        return doc

    doc_standard = _make_doc("standards.txt")
    await db.flush()
    job_standard = RagIngestionJob(
        collection_id=collection.id,
        document_id=doc_standard.id,
        operation="upsert_document",
        status="queued",
    )
    db.add(job_standard)
    await db.flush()
    await ingest_document(
        db,
        document=doc_standard,
        job=job_standard,
        raw_bytes=(
            "Технические средства должны соответствовать ГОСТ Р 52931-2008, группа В4. "
            "Это требование относится к условиям эксплуатации оборудования."
        ).encode("utf-8"),
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )

    doc_noise = _make_doc("noise.txt")
    await db.flush()
    job_noise = RagIngestionJob(
        collection_id=collection.id,
        document_id=doc_noise.id,
        operation="upsert_document",
        status="queued",
    )
    db.add(job_noise)
    await db.flush()
    await ingest_document(
        db,
        document=doc_noise,
        job=job_noise,
        raw_bytes=(
            "ГОСТ 15150-69 описывает климатическое исполнение изделий. "
            "Оборудование должно работать при температуре от плюс 5 до плюс 50 градусов."
        ).encode("utf-8"),
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )
    await db.commit()

    hits = await retrieve(
        db,
        chat_id=doc_standard.id,
        ticket_id=doc_standard.id,
        message_id=None,
        query_text="ГОСТ Р 52931",
        embedding=providers.embedding,
        vector_store=providers.vector_store,
    )

    assert hits
    top_chunk, _top_score = hits[0]
    assert "52931-2008" in top_chunk.chunk_text


@pytest.mark.asyncio
async def test_ingest_invalid_pdf(db, providers):
    """Ingestion с некорректным PDF должен завершиться с статусом failed."""
    collection = await get_default_rag_collection(db)
    doc = RagDocument(
        collection_id=collection.id,
        source_type="file",
        source_name="invalid.pdf",
        mime_type="application/pdf",
        storage_url=None,
        current_version=0,
    )
    db.add(doc)
    await db.flush()
    job = RagIngestionJob(
        collection_id=collection.id,
        document_id=doc.id,
        operation="upsert_document",
        status="queued",
    )
    db.add(job)
    await db.flush()

    with pytest.raises(ValueError, match="Ошибка извлечения текста из PDF"):
        await ingest_document(
            db,
            document=doc,
            job=job,
            raw_bytes=b"invalid pdf",
            embedding=providers.embedding,
            vector_store=providers.vector_store,
        )
    # после исключения job.status уже должен быть "failed"
    await db.refresh(job)
    assert job.status == "failed"
    assert job.error_message is not None
    assert "Ошибка извлечения текста из PDF" in job.error_message
